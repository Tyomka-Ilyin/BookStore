import mimetypes
import os

import docx
from django.contrib import messages
from django.http import HttpResponse, HttpResponseRedirect
from django.shortcuts import render
from django.urls import reverse

from Buy.models import Buy
from .forms import ReportForm


def runReport(request):
    if request.method == 'POST':
        # create a form instance and populate it with data from the request:
        form = ReportForm(request.POST)
        # check whether it's valid:
        if form.is_valid():
            dateStart = form.cleaned_data['dateStart']
            dateEnd = form.cleaned_data['dateEnd']

            buys = Buy.objects.filter(dateTime__gte=dateStart, dateTime__lte=dateEnd)

            sum = 0

            mymessage=messages

            for buy in buys:
                sum+=buy.sumBuy

            if sum == 0:
                messages.error(request, "В данном периоде нет покупок")
                return HttpResponseRedirect('/admin', {"Mymessages": messages})
            else:

                doc = docx.Document()
                doc.add_heading("Умная книга")

                doc.add_paragraph("Выручка в период с "+f"{dateStart}"+" по "+f"{dateEnd}"+" составляет :"+f"{sum}"+" руб.")

                table = doc.add_table(rows=len(buys), cols=4)
                # применяем стиль для таблицы
                table.style = 'Table Grid'

                for row in range(len(buys)):
                    str = buys[row]
                    cell = table.cell(row - 1, 0)
                    cell.text = f"Номер покупки: {str.id}"

                for row in range(len(buys)):
                    str = buys[row]
                    cell = table.cell(row - 1, 1)
                    cell.text = f"Покупатель: {str.user.username}"

                for row in range(len(buys)):
                    str = buys[row]
                    cell = table.cell(row - 1, 2)
                    cell.text = f"Дата: {str.dateTime}"

                for row in range(len(buys)):
                    str = buys[row]
                    cell = table.cell(row - 1, 3)
                    cell.text = f"Покупатель: {str.sumBuy} руб."

                doc.save("media/reports/Отчет по выручке.docx")

                filename = 'media/reports/Отчет по выручке.docx'
                data = open(filename, "rb").read()
                response = HttpResponse(data, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                response['Content-Disposition'] = 'attachment; filename=Otchet po virychke.docx'
                response['Content-Length'] = os.path.getsize(filename)

                return response

        else:
            messages.error(request, "Заполните все поля")
            return HttpResponseRedirect('/admin', {"Mymessages": messages})