from django.contrib import messages
from django.core.mail import send_mail, EmailMessage
from django.db import transaction
from django.http import HttpResponse, HttpResponseRedirect, JsonResponse
from django.shortcuts import render
from django.views.generic.base import View

from .forms import BuyForm
from .models import Buy, structureBuy, discount
from Cart.models import Cart, structureCart
from DetailUser.models import Bonus

import datetime

import math

import docx

class BuyView(View):
    def get(self, request, *args, **kwargs):
        form = BuyForm(request.POST or None)

        cart = Cart.objects.get(user=request.user)

        structure = structureCart.objects.filter(cart=cart)

        bonus = Bonus.objects.get(user=request.user)

        return render(request, 'buy/buy.html', {"structure": structure, "cart": cart, "bonus": bonus, "form": form})

    def post(self, request, *args, **kwargs):
        form = BuyForm(request.POST or None)

        cart = Cart.objects.get(user=request.user)
        structure = structureCart.objects.filter(cart=cart)
        bonusUser = Bonus.objects.get(user=request.user)

        context = {'form': form, "structure": structure, "cart": cart, "bonus": bonusUser}

        if form.is_valid():
            bonusForm = form.cleaned_data['bonus']

            doc = docx.Document()
            title = doc.add_heading("Умная книга")

            if (bonusForm == ''):
                sumBuy = cart.sumCart

                bonusUser = Bonus.objects.get(user=request.user)
                bonusUser.kolBonus = bonusUser.kolBonus + math.ceil(sumBuy / 100 * 10)
                bonusUser.save()

                buy = Buy(user=request.user, dateTime=datetime.datetime.now().strftime("%Y-%m-%d %H:%M"),
                          sumBuy=sumBuy)
                buy.save()

                doc.add_paragraph(f"Чек по покупке № {buy.id}")
                doc.add_paragraph(f"Дата покупки: {buy.dateTime}")

                buys = []

                for str in structure:
                    structurebuy = structureBuy(buy=buy, book=str.book)
                    structurebuy.save()
                    buys.append(structurebuy)

                table = doc.add_table(rows=len(buys), cols=2)
                # применяем стиль для таблицы
                table.style = 'Table Grid'

                for row in range(len(buys)):
                    str = buys[row-1]
                    cell = table.cell(row-1, 0)
                    cell.text = f"Название: {str.book.title}"

                for row in range(len(buys)):
                    str = buys[row-1]
                    cell = table.cell(row-1, 1)
                    cell.text = f"Цена: {str.book.nds.priceWithNDS} руб."

                doc.add_heading(f"Сумма покупки: {sumBuy} руб.")

                doc.save(f"media/checks/Чек покупки №{buy.id}.docx")

                books = []

                for str in structure:
                    books.append(str.book.bookFile.path)

                email = EmailMessage(
                    'Сообщение от онлайн-магазина "Умная книга"', 'Спасибо за покупку в нашем магазине!',
                    'bookstore2000@mail.ru', [request.user.email])

                for book in books:
                    email.attach_file(book)

                email.attach_file(f"media/checks/Чек покупки №{buy.id}.docx")

                email.send()

                structure.delete()
                cart.sumCart = 0
                cart.save()

                return HttpResponseRedirect("/")
            else:
                if (int(bonusForm) <= int(Bonus.objects.get(user=request.user).kolBonus)):
                    sumBuy = int(cart.sumCart) - int(bonusForm)
                    if (math.ceil(sumBuy / 100 * 40) >= int(bonusForm)):
                        buy = Buy(user=request.user, dateTime=datetime.datetime.now().strftime("%Y-%m-%d %H:%M"),
                                  sumBuy=sumBuy)
                        buy.save()

                        doc.add_paragraph(f"Чек по покупке № {buy.id}")
                        doc.add_paragraph(f"Дата покупки: {buy.dateTime}")

                        bonusUser = Bonus.objects.get(user=request.user)
                        bonusUser.kolBonus = int(bonusUser.kolBonus) - int(bonusForm)
                        bonusUser.save()

                        buys = []

                        for str in structure:
                            structurebuy = structureBuy(buy=buy, book=str.book)
                            structurebuy.save()
                            buys.append(structurebuy)

                        table = doc.add_table(rows=len(buys), cols=2)
                        # применяем стиль для таблицы
                        table.style = 'Table Grid'

                        for row in range(len(buys)):
                            str = buys[row - 1]
                            cell = table.cell(row - 1, 0)
                            cell.text = f"Название: {str.book.title}"

                        for row in range(len(buys)):
                            str = buys[row - 1]
                            cell = table.cell(row - 1, 1)
                            cell.text = f"Цена: {str.book.nds.priceWithNDS} руб."

                        disc = discount(buy=buy, amounDiscount=bonusForm)
                        disc.save()

                        doc.add_heading(f"Сумма покупки: {sumBuy} руб.")
                        doc.add_heading(f"Скидка: {bonusForm} руб.")

                        doc.save(f"media/checks/Чек покупки №{buy.id}.docx")

                        books = []

                        for str in structure:
                            books.append(str.book.bookFile.path)

                        email = EmailMessage(
                            'Сообщение от онлайн-магазина "Умная книга"', 'Спасибо за покупку в нашем магазине!',
                            'bookstore2000@mail.ru', [request.user.email])

                        for book in books:
                            email.attach_file(book)

                        email.attach_file(f"media/checks/Чек покупки №{buy.id}.docx")

                        email.send()

                        structure.delete()
                        cart.sumCart = 0
                        cart.save()

                        return HttpResponseRedirect("/")
                    else:
                        form.add_error('bonus', 'Бонусами можно оплатить только до 40% суммы покупки')
                else:
                    form.add_error('bonus', "Не достаточно бонусов")
                    return render(request, 'buy/buy.html', context)

        return render(request, 'buy/buy.html', context)


def goBuy(request):
    cart = Cart.objects.get(user=request.user)
    structure = structureCart.objects.filter(cart=cart)
    bonusUser = Bonus.objects.get(user=request.user)

    bonus = request.POST["bonus"]

    buy = Buy(user=request.user, dateTime=datetime.datetime.now(), sumBuy=int(cart.sumCart) - int(bonus))
    buy.save()

    for str in structure:
        structurebuy = structureBuy(buy=buy, book=str.book)
        structurebuy.save()

    if (bonus != 0):
        disc = discount(buy=buy, amounDiscount=bonus)
        disc.save()

    books = []

    for str in structure:
        books.append(str.book.bookFile.path)

    email = EmailMessage(
        'Сообщение от онлайн-магазина "Умная книга"', 'Спасибо за покупку в нашем магазине!', 'bookstore2000@mail.ru',
        [request.user.email])

    for book in books:
        email.attach_file(book)

    email.send()

    structure.delete()
    cart.sumCart = 0
    cart.save()

    return HttpResponseRedirect("/")
